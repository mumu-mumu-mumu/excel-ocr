#encoding:utf-8
#mumu

from logging import RootLogger
import requests
import base64
import pandas as pd
import os
import json
import configparser
from tkinter import *
from tkinter import filedialog
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook

import openpyxl





# 检查配置文件是否存在，如果不存在就创建一个
if not os.path.exists('config.ini'):
    config = configparser.ConfigParser()
    config['DEFAULT'] = {'API_Key': '', 'Secret_Key': ''}
    with open('config.ini', 'w') as f:
        config.write(f)

# 定义一个选择文件函数
def select_folder():
    global folder_path
    folder_path = filedialog.askdirectory()
    if folder_path != '':
        print("您选择的文件夹路径为：", folder_path)
    else:
        print("您没有选择文件夹路径")

# 定义一个保存用户输入的API Key和Secret Key函数
def save_keys():
    
    API_Key = api_key_entry.get()
    Secret_Key = secret_key_entry.get()
    config.set('DEFAULT', 'API_Key', API_Key)
    config.set('DEFAULT', 'Secret_Key', Secret_Key)
    with open('config.ini', 'w') as f:
        config.write(f)





#定义智能财务票据识别程序
def finance_url():
    
    request_url = "https://aip.baidubce.com/rest/2.0/ocr/v1/multiple_invoice"
    url = "https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id=%s&client_secret=%s"%(API_Key,Secret_Key)

    payload = ""
    headers = {
        'Content-Type': 'application/json',
        'Accept': 'application/json'
        }
        
    response = requests.request("POST", url, headers=headers, data=payload)
    response_str = response.content.decode('utf-8')
    response_dict = json.loads(response_str)
    access_token = response_dict['access_token']
    print(access_token)

    #搜索文件
    input_folder = folder_path
    os.chdir(input_folder)
    catalog = os.listdir()
    for each in catalog:
        if each.endswith('.jpg') or each.endswith('.jpeg') or each.endswith('.png'):
            # 二进制方式打开图片文件
            f = open(each, 'rb')
            img = base64.b64encode(f.read())
            print(img)
            params = {"image":img}
            request_url = request_url + "?access_token=" + access_token
            headers = {'content-type': 'application/x-www-form-urlencoded'}
            response = requests.post(request_url, data=params, headers=headers)
            if response:
                print(response.json())
                print(response.text)

            #把文件以json格式保存起来
            name = each.split('.')[0]
            with open(f'{name}.json', 'w', encoding = 'utf-8') as z:
                z.write(response.text)

    input_folder = folder_path
    output_folder = "%s\\excel" % folder_path

    # 如果output路径不存在则自动生成excel文件夹
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    # 遍历文件夹中的JSON文件
    for filename in os.listdir(input_folder):
        if filename.endswith(".json"):
            # 读取JSON文件内容并转换为Python对象
            with open(os.path.join(input_folder, filename), "r", encoding='utf-8') as f:
                data = json.load(f)
                result = data['words_result'][0]['result']
            
            def get_value(result, key):
                # 如果result中有key对应的值，就返回它的第一个元素的'word'属性
                if result.get(key):
                    return result[key][0]['word']
                # 否则就返回None
                else:
                    return None
            
            keys = [
                "AmountInWords", "InvoiceNumConfirm", "CommodityEndDate", "CommodityEndDate", "CommodityVehicleType",
                "CommodityPrice", "InvoiceTag", "NoteDrawer", "SellerAddress", "CommodityNum", "SellerRegisterNum",
                "CommodityTaxRate", "ServiceType", "TotalTax", "InvoiceCodeConfirm", "CheckCode", "InvoiceCode",
                "InvoiceDate", "PurchaserRegisterNum", "InvoiceTypeOrg", "Password", "OnlinePay", "Agent",
                "AmountInFiguers", "PurchaserBank", "Checker", "City", "TotalAmount", "CommodityAmount", "PurchaserName",
                "CommodityType", "Province", "InvoiceType", "SheetNum", "PurchaserAddress", "CommodityTax",
                "CommodityPlateNum", "CommodityUnit", "Payee", "CommodityName", "SellerName", "InvoiceNum",
                'Invoice code', 'Invoice number', 'Invoice date', 'Seller name', 'Seller address',
                'Seller register number', 'date', 'starting_station', 'Waiting_area', 'ticket_num',
                'train_num', 'ticket_rates', 'serial_number', 'ID_card', 'seat_category', 'ServiceType',
                'destination_station', 'name', 'sales_station', 'time', 'seat_num', "Remarks", "SellerBank", "MachineCode"
            ]
            
            for key in keys:
                # 调用get_value函数，获取值
                value = get_value(result, key)
                # 把键和值放到字典中
                data[key] = [value]
            
            df = pd.DataFrame(data)
            # 删除空列
            df = df.dropna(axis=1)

            # 保存为Excel文件
            output_file = os.path.splitext(filename)[0] + ".xlsx"
            output_path = os.path.join(output_folder, output_file)
            df.to_excel(output_path, index=False)
            replace_dict = {"AmountInWords": "金额大写", "InvoiceNumConfirm": "发票号码确认", "CommodityEndDate": "商品截止日期", 
                                'date':'日期','starting_station':"上车站",'Waiting_area':'候车区','ticket_num':'车票号','train_num':'车次',
                                "ticket_rates":'票价','serial_number':'',"ID_card":'证件号','seat_category':'席别',
                                'destination_station':'下车站','name':'姓名','sales_sation':'售票站','time':'时间','seat_num':'座位号',}
            # 使用 replace 方法，把 DataFrame 中的值按照字典进行替换
            
            df = df.rename(columns=replace_dict)
            print('%s成功完成转换' % filename)
    input_folder_excel = os.path.join(input_folder, "excel")
    os.chdir(input_folder_excel)
    catalog = os.listdir()
    for each in catalog:
        if each.endswith('.xlsx'):
            excel_file = os.path.join(input_folder_excel, each)
            df = pd.read_excel(excel_file, engine='openpyxl')
            replace_dict = {"AmountInWords": "金额大写", "InvoiceNumConfirm": "发票号码确认", "CommodityEndDate": "商品截止日期", 
                            'date':'日期','starting_station':"上车站",'Waiting_area':'候车区','ticket_num':'车票号','train_num':'车次',
                            "ticket_rates":'票价','serial_number':'',"ID_card":'证件号','seat_category':'席别',
                            'destination_station':'下车站','name':'姓名','sales_sation':'售票站','time':'时间','seat_num':'座位号',}
        # 使用 replace 方法，把 DataFrame 中的值按照字典进行替换
        df = df.replace(replace_dict)
        df = df.rename(columns=replace_dict)
        
        # 生成输出文件路径
        file_name, extension = os.path.splitext(each)
        output_file = file_name  + extension
        output_path = os.path.join(output_folder, output_file)
        
        # 保存为新的 Excel 文件
        df.to_excel(output_path, index=False)

    last_label = Label(root, text='已完成全部转换', bg='lightyellow')
    last_label.grid(row=6, column=0, padx=5, pady=5)


#定义手写文字识别url
def handwriting_url():
    
    request_url = "https://aip.baidubce.com/rest/2.0/ocr/v1/handwriting"
    url = "https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id=%s&client_secret=%s"%(API_Key,Secret_Key)

    payload = ""
    headers = {
        'Content-Type': 'application/json',
        'Accept': 'application/json'
        }
        
    response = requests.request("POST", url, headers=headers, data=payload)
    response_str = response.content.decode('utf-8')
    response_dict = json.loads(response_str)
    access_token = response_dict['access_token']
    print(access_token)

    #搜索文件
    input_folder = folder_path
    os.chdir(input_folder)
    catalog = os.listdir()
    for each in catalog:
        if each.endswith('.jpg') or each.endswith('.jpeg') or each.endswith('.png'):
            # 二进制方式打开图片文件
            f = open(each, 'rb')
            img = base64.b64encode(f.read())
            print(img)
            params = {"image":img}
            request_url = request_url + "?access_token=" + access_token
            headers = {'content-type': 'application/x-www-form-urlencoded'}
            response = requests.post(request_url, data=params, headers=headers)
            if response:
                print(response.json())
                print(response.text)

            #把文件以json格式保存起来
            name = each.split('.')[0]
            with open(f'{name}.json', 'w', encoding = 'utf-8') as z:
                z.write(response.text)

    input_folder = folder_path
    output_folder = "%s\\word"%(folder_path)
    #如果无output路径不存在则自动生成excel文件夹
    file_path = output_folder
    if os.path.exists(file_path) is False:
        os.makedirs(file_path)
    # 遍历文件夹中的 JSON 文件
    for filename in os.listdir(input_folder):
        if filename.endswith(".json"):
            # 读取 JSON 文件内容并转换为 Python 对象
            with open(os.path.join(input_folder, filename), "r", encoding='utf-8') as f:
                data = json.load(f)
                document = Document()
                 # 设置字体
            document.styles['Normal'].font.name = '宋体'
            document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            document.styles['Normal'].font.size = Pt(12)

            # 将JSON数据中的文本写入Word文档
            for item in data['words_result']:
                text = item['words']
                paragraph = document.add_paragraph(text)
    # 保存为 Docx 文件
            output_file = os.path.splitext(filename)[0] + ".docx"
            output_path = os.path.join(output_folder, output_file)
            document.save(output_path)
            print('%s成功完成转换'%(filename))
            

#处理成功
    last_label = Label(root,text='已完成全部转换',bg='lightyellow')
    last_label.grid(row=6,column=0,padx=5,pady=5)





#定义通用文字识别url
def text_url():
    
    request_url = "https://aip.baidubce.com/rest/2.0/ocr/v1/accurate_basic"
    url = "https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id=%s&client_secret=%s"%(API_Key,Secret_Key)

    payload = ""
    headers = {
    'Content-Type': 'application/json',
    'Accept': 'application/json'
    }

    response = requests.request("POST", url, headers=headers, data=payload)
    response_str = response.content.decode('utf-8')
    response_dict = json.loads(response_str)
    access_token = response_dict['access_token']
    print(access_token)

    #搜索文件
    input_folder = folder_path
    os.chdir(input_folder)
    catalog = os.listdir()
    for each in catalog:
        if each.endswith('.jpg') or each.endswith('.jpeg') or each.endswith('.png'):
            # 二进制方式打开图片文件
            f = open(each, 'rb')
            img = base64.b64encode(f.read())
            print(img)
            params = {"image":img}
            request_url = request_url + "?access_token=" + access_token
            headers = {'content-type': 'application/x-www-form-urlencoded'}
            response = requests.post(request_url, data=params, headers=headers)
            if response:
                print(response.json())
                print(response.text)

            #把文件以json格式保存起来
            name = each.split('.')[0]
            with open(f'{name}.json', 'w', encoding = 'utf-8') as z:
                z.write(response.text)

    input_folder = folder_path
    output_folder = "%s\\docx"%(folder_path)
    #如果无output路径不存在则自动生成docx文件夹
    file_path = output_folder
    if os.path.exists(file_path) is False:
        os.makedirs(file_path)
    # 遍历文件夹中的 JSON 文件
    for filename in os.listdir(input_folder):
        if filename.endswith(".json"):
            # 读取 JSON 文件内容并转换为 Python 对象
            with open(os.path.join(input_folder, filename), "r", encoding='utf-8') as f:
                data = json.load(f)
                document = Document()
            # 设置字体
            document.styles['Normal'].font.name = '宋体'
            document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            document.styles['Normal'].font.size = Pt(12)

            # 将JSON数据中的文本写入Word文档
            for item in data['words_result']:
                text = item['words']
                paragraph = document.add_paragraph(text)
            # 保存为 Docx 文件
            output_file = os.path.splitext(filename)[0] + ".docx"
            output_path = os.path.join(output_folder, output_file)
            document.save(output_path)
            print('%s成功完成转换'%(filename))
    #处理成功
    last_label = Label(root,text='已完成全部转换',bg='lightyellow')
    last_label.grid(row=6,column=0,padx=5,pady=5)





#定义表格识别模块
def table_url():
    
    request_url = "https://aip.baidubce.com/rest/2.0/ocr/v1/table"

    url = "https://aip.baidubce.com/oauth/2.0/token?grant_type=client_credentials&client_id=%s&client_secret=%s"%(API_Key,Secret_Key)

    payload = ""
    headers = {
        'Content-Type': 'application/json',
        'Accept': 'application/json'
        }
        
    response = requests.request("POST", url, headers=headers, data=payload)
    response_str = response.content.decode('utf-8')
    response_dict = json.loads(response_str)
    access_token = response_dict['access_token']
    print(access_token)

    #搜索文件
    input_folder = folder_path
    os.chdir(input_folder)
    catalog = os.listdir()
    for each in catalog:
        if each.endswith('.jpg') or each.endswith('.jpeg') or each.endswith('.png'):
            # 二进制方式打开图片文件
            f = open(each, 'rb')
            img = base64.b64encode(f.read())
            print(img)
            params = {"image":img}
            request_url = request_url + "?access_token=" + access_token
            headers = {'content-type': 'application/x-www-form-urlencoded'}
            response = requests.post(request_url, data=params, headers=headers)
            if response:
                print(response.json())
                print(response.text)

            #把文件以json格式保存起来
            name = each.split('.')[0]
            with open(f'{name}.json', 'w', encoding = 'utf-8') as z:
                z.write(response.text)

    input_folder = folder_path
    output_folder = "%s\\excel"%(folder_path)
    #如果无output路径不存在则自动生成excel文件夹
    file_path = output_folder
    if os.path.exists(file_path) is False:
        os.makedirs(file_path)
    # 遍历文件夹中的 JSON 文件
    for filename in os.listdir(input_folder):
        if filename.endswith(".json"):
            # 读取 JSON 文件内容并转换为 Python 对象
            with open(os.path.join(input_folder, filename), "r", encoding='utf-8') as f:
                data = json.load(f)
            table_data = data['tables_result'][0]['body']
            # 将表格数据转换成pandas DataFrame对象
            df = pd.DataFrame()
            for row in table_data:
                # 获取单元格位置和内容
                row_start = row['row_start']
                row_end = row['row_end']
                col_start = row['col_start']
                col_end = row['col_end']
                words = row['words']
                # 将单元格内容添加到DataFrame中
                for r in range(row_start, row_end+1):
                    for c in range(col_start, col_end+1):
                        df.loc[r, c] = words

            # 保存为 Excel 文件
            output_file = os.path.splitext(filename)[0] + ".xlsx"
            output_path = os.path.join(output_folder, output_file)
            df.to_excel(output_path, index=False)
            print('%s成功完成转换'%(filename))
    
    #处理成功
    last_label = Label(root,text='已完成全部转换',bg='lightyellow')
    last_label.grid(row=6,column=0,padx=5,pady=5)

def merged_excel() :
    folder_path_excel = '%s\\excel'%folder_path

    # 创建一个新的Excel文件
    merged_wb = Workbook()

    # 遍历要合并的每个Excel文件
    for file_name in os.listdir(folder_path_excel):
        # 加载当前Excel文件
        if file_name.endswith(".xlsx"):
            file_path = os.path.join(folder_path_excel, file_name)
            new_sheet_name = file_name.split('.')[0]
            wb = load_workbook(file_path)
            # 遍历每个sheet页
            for sheet_name in wb.sheetnames:
                # 读取当前sheet页中的所有行和列数据
                sheet = wb[sheet_name]
                rows = sheet.rows
                cols = sheet.columns
                # 创建一个新的sheet页
                merged_sheet = merged_wb.create_sheet(title=f"{new_sheet_name}")
                # 将当前sheet页中的所有数据复制到新的sheet页中
                for row_idx, row in enumerate(rows, 1):
                    for col_idx, cell in enumerate(row, 1):
                        merged_sheet.cell(row=row_idx, column=col_idx, value=cell.value)

    # 保存合并后的Excel文件
    merged_file_path = os.path.join(folder_path_excel, "merged.xlsx")
    merged_wb.save(merged_file_path)  
    #处理成功
    last_label = Label(root,text='已完成整合',bg='lightyellow')
    last_label.grid(row=6,column=0,padx=5,pady=5)

#def merged_word() :


# 创建一个窗口
root = Tk()
root.title("OCR小工具")
root.geometry("800x400")
root.resizable(True, True)

# 创建两个标签和两个输入框
api_key_label = Label(root, text="API Key:",)
api_key_entry = Entry(root)
secret_key_label = Label(root, text="Secret Key:")
secret_key_entry = Entry(root)
tips_label = Label(root,text='目前仅支持.jpg、.jpeg、.png的格式图片')

#创建ocr类型按钮
finance_button = Button(root, text="财务票据", command=finance_url)
handwriting_button = Button(root, text="手写文件", command=handwriting_url)
text_button = Button(root, text="普通文件", command=text_url)
teble_button = Button(root, text="表格文件", command=table_url)

#创建文件整合按钮
merged_excel_button = Button(root, text='excel整合', command=merged_excel)
#merged_word_button = Button(root, text='word整合', command=merged_word)

# 读取配置文件，如果已经有保存的API Key和Secret Key，就显示在输入框中
config = configparser.ConfigParser()
config.read('config.ini')
API_Key = config.get('DEFAULT', 'API_Key')
Secret_Key = config.get('DEFAULT', 'Secret_Key')
if API_Key:
    api_key_entry.insert(0, API_Key)
if Secret_Key:
    secret_key_entry.insert(0, Secret_Key)

# 创建保存按钮和文件夹选择按钮
save_button = Button(root, text="保存", command=save_keys)
folder_button = Button(root, text="选择文件夹", command=lambda: select_folder())

# 使用网格布局来排列标签、输入框和按钮
api_key_label.grid(row=0, column=0, padx=5, pady=5, sticky=E)
api_key_entry.grid(row=0, column=1, padx=5, pady=5, sticky=W)
secret_key_label.grid(row=1, column=0, padx=5, pady=5, sticky=E, )
secret_key_entry.grid(row=1, column=1, padx=5, pady=5, sticky=W)
save_button.grid(row=2, column=0, padx=5, pady=5)
folder_button.grid(row=2, column=1, padx=5, pady=5)
tips_label.grid(row=2,column=3, columnspan=2,padx=5, pady=5)
finance_button.grid(row=4,column=0, padx=5, pady=5)
handwriting_button.grid(row=4,column=1, padx=5, pady=5)
text_button.grid(row=4,column=2, padx=5, pady=5)
teble_button.grid(row=4,column=3, padx=5, pady=5)
merged_excel_button.grid(row=5,column=0,padx=5,pady=5)
#merged_word_button.grid(row=5,column=1,padx=5,pady=5)

# 进入消息循环，等待用户输入
root.mainloop()
#software was created by mumu




